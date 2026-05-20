import streamlit as st
import requests
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import base64

# --- 1. إعدادات الهوية البصرية وتنسيق RTL ---
st.set_page_config(page_title="Innovation Radar Pro v2", layout="wide")
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@400;500;700&display=swap');
    
    /* إجبار التطبيق بالكامل على اتجاه اليمين والمحاذاة لليمين */
    html, body, [data-testid="stAppViewContainer"], .stApp {
        direction: rtl !important;
        text-align: right !important;
        font-family: 'Tajawal', sans-serif !important;
    }
    
    /* منع التوسيط في المتصفح للحفاظ على القراءة اليمينية المريحة */
    p, li, span, div, label {
        direction: rtl !important;
        text-align: right !important;
        line-height: 1.6 !important;
        display: block; 
    }

    /* ضمان محاذاة العناوين لليمين في المتصفح */
    h1, h2, h3, h4, h5, h6 {
        text-align: right !important;
        color: #1e3a8a !important;
        direction: rtl !important;
    }

    /* تحسين شكل صناديق التنبيه لتكون محاذاتها يميناً */
    div[data-testid="stMarkdownContainer"] > div {
        text-align: right !important;
    }
    </style>
""", unsafe_allow_html=True)


# ==========================================
# 2. دالة الاتصال بالمحرك (جوجل API)
# ==========================================
def call_pro_api(prompt, image_file=None):
    if "PRO_API_KEY" not in st.secrets:
        return "خطأ: مفتاح API غير موجود في إعدادات Secrets"
        
    api_key = st.secrets["PRO_API_KEY"]
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-flash-latest:generateContent?key={api_key}"
    headers = {'Content-Type': 'application/json'}
    
    try:
        if image_file:
            img_data = base64.b64encode(image_file.read()).decode('utf-8')
            payload = {
                "contents": [{
                    "parts": [
                        {"text": prompt},
                        {"inline_data": {"mime_type": "image/jpeg", "data": img_data}}
                    ]
                }]
            }
        else:
            payload = {
                "contents": [{"parts": [{"text": prompt}]}]
            }

        response = requests.post(url, json=payload, headers=headers, timeout=30)
        
        if response.status_code == 200:
            return response.json()['candidates'][0]['content']['parts'][0]['text']
        else:
            return f"خطأ في الاتصال: {response.status_code} - {response.text}"
            
    except Exception as e:
        return f"فشل النظام في الاستجابة: {str(e)}"


# ==========================================
# 3. دالة توليد ملف الوورد المنسق (التوسيط الكامل بناءً على رغبتك)
# ==========================================
def create_docx(report_text, idea_title):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        import re
        from io import BytesIO

        doc = Document()

        # إجبار المستند بنائياً على اتجاه اليمين (RTL Document Structure)
        section = doc.sections[0]
        sectPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        sectPr.append(bidi)

        # استخراج الرقم برمجياً قبل تنظيف النص لتنسيقه بشكل مخصص
        score_match = re.search(r'\[===SCORE===\]\s*(\d+)\s*\[===/SCORE===\]', report_text)
        innovation_score = score_match.group(1) if score_match else None

        # تنظيف النص تماماً وبشكل صارم من كل الأوسمة والنقاط الزائدة لملف الوورد
        clean_text = re.sub(r'\[===SCORE===\].*?\[===/SCORE===\]', '', report_text, flags=re.DOTALL)
        clean_text = re.sub(r'\[===.*?===\]', '', clean_text)
        clean_text = clean_text.replace('###', '').replace('---', '').replace('**', '').replace('*', '')

        # استخلاص سطر واحد مختصر كعنوان للابتكار
        short_title = idea_title.split('\n')[0].strip()
        if len(short_title) > 60:
            short_title = short_title[:60] + "..."

        # تعريف الألوان الرسمية للتقرير
        MAIN_COLOR = RGBColor(30, 58, 138)  # أزرق داكن
        TEXT_COLOR = RGBColor(0, 0, 0)      # أسود للمتن

        # --- أ. العنوان الرئيسي في الأعلى (موسط) ---
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.right_to_left = True
        
        run_title = p_title.add_run('تقرير رادار الابتكار الاحترافي') [cite: 1]
        run_title.bold = True
        run_title.font.size = Pt(18)
        run_title.font.name = 'Arial'
        run_title.font.color.rgb = MAIN_COLOR
        run_title._element.get_or_add_rPr().get_or_add_rtl().val = True

        # --- ب. عنوان الابتكار المخصص (موسط) ---
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.right_to_left = True
        
        run_sub = p_sub.add_run(f"عنوان الابتكار: {short_title}") [cite: 2]
        run_sub.font.size = Pt(14)
        run_sub.font.name = 'Arial'
        run_sub.font.color.rgb = RGBColor(100, 116, 139)
        run_sub._element.get_or_add_rPr().get_or_add_rtl().val = True

        # --- ج. إضافة مؤشر الفرادة بشكل رسمي مخصص (موسط) ---
        if innovation_score:
            p_score = doc.add_paragraph()
            p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_score.paragraph_format.right_to_left = True
            p_score.paragraph_format.space_before = Pt(6)
            p_score.paragraph_format.space_after = Pt(24)
            
            score_num = int(innovation_score)
            SCORE_COLOR = RGBColor(16, 185, 129) if score_num >= 85 else (RGBColor(245, 158, 11) if score_num >= 65 else RGBColor(239, 68, 68))
            
            run_score_label = p_score.add_run("مؤشر الفرادة المحتملة للابتكار: ")
            run_score_label.font.size = Pt(13)
            run_score_label.font.name = 'Arial'
            run_score_label.font.color.rgb = RGBColor(71, 85, 105)
            run_score_label._element.get_or_add_rPr().get_or_add_rtl().val = True
            
            run_score_val = p_score.add_run(f"{innovation_score}%")
            run_score_val.bold = True
            run_score_val.font.size = Pt(16)
            run_score_val.font.name = 'Arial'
            run_score_val.font.color.rgb = SCORE_COLOR
            run_score_val._element.get_or_add_rPr().get_or_add_rtl().val = True
        else:
            p_sub.paragraph_format.space_after = Pt(24)

        # --- د. متن التقرير (التوسيط الكامل بناءً على التحديث الصارم) ---
        for para in clean_text.split('\n'):
            text = para.strip()
            if text:
                p = doc.add_paragraph()
                p.paragraph_format.right_to_left = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # جعل كافة الفقرات موسطة تماماً في الصفحة
                p.paragraph_format.line_spacing = 1.3
                
                run = p.add_run(text)
                run.font.name = 'Arial'
                
                rPr = run._element.get_or_add_rPr()
                rPr.get_or_add_rtl().val = True
                
                keywords = ["التشخيص", "المطالبات", "الجدوى", "الفرادة", "توصية", "المنافسون", "خارطة"]
                if any(h in text for h in keywords):
                    run.bold = True
                    run.font.size = Pt(16)
                    run.font.color.rgb = MAIN_COLOR
                    p.paragraph_format.space_before = Pt(14)
                    p.paragraph_format.space_after = Pt(6)
                else:
                    run.font.size = Pt(14)
                    run.font.color.rgb = TEXT_COLOR
                    p.paragraph_format.space_after = Pt(8)

        # --- هـ. حقوق الإعداد والتطوير في الأسفل (موسطة أيضاً) ---
        doc.add_paragraph() 
        p_footer = doc.add_paragraph()
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_footer.paragraph_format.right_to_left = True
        
        run_line = p_footer.add_run("________________________________")
        doc.add_paragraph()
        
        p_credit = doc.add_paragraph()
        p_credit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_credit.paragraph_format.right_to_left = True
        
        run_credit = p_credit.add_run('إعداد وتطوير: أ. خالد العقبي | المسار الرقمي') [cite: 57]
        run_credit.bold = True
        run_credit.font.size = Pt(11)
        run_credit.font.name = 'Arial'
        run_credit.font.color.rgb = MAIN_COLOR
        run_credit._element.get_or_add_rPr().get_or_add_rtl().val = True

        bio = BytesIO()
        doc.save(bio)
        return bio.getvalue()
        
    except Exception as e:
        st.error(f"حدث خطأ أثناء توليد ملف Word: {str(e)}")
        return None


# ==========================================
# 4. واجهة المستخدم (التوجيه الصارم، حماية الحصص والروابط الحية)
# ==========================================
st.markdown("<h1 style='text-align:center;'>🛡️ رادار الابتكار Pro</h1>", unsafe_allow_html=True)

if "gate_passed" not in st.session_state:
    st.session_state.gate_passed = False
if "full_report" not in st.session_state:
    st.session_state.full_report = None

if not st.session_state.gate_passed:
    col1, col2 = st.columns([2, 1])
    with col1:
        st.markdown("### 📋 استمارة الفحص البنائي للابتكار")
        input_1 = st.text_input("✍️ 1. اسم الابتكار ووظيفته الأساسية:", placeholder="مثال: خاتم ذكي لحساب السعرات الحرارية...")
        input_2 = st.text_area("🔧 2. الآلية التقنية (الخلطة السرية):", height=100, placeholder="كيف يعمل تقنياً؟ اذكر الحساسات، الخوارزميات، أو طريقة معالجة البيانات...")
        input_3 = st.text_area("🎯 3. المشكلة والفجوة في السوق:", height=100, placeholder="ما هي المشكلة التي يحلها؟ وما الذي يميّزه عن الساعات الذكية أو الحلول البديلة الحالية؟")
    with col2:
        uploaded_file = st.file_uploader("🖼️ ارفع رسم كروكي أو مخطط هندسي (اختياري):", type=["jpg", "png", "jpeg"])
    
    if st.button("بدء الفحص الاستراتيجي 🚀"):
        if input_1.strip() and input_2.strip() and input_3.strip():
            with st.spinner("جاري تحليل البيانات والصور وتطبيق بروتوكول الريادة الابتكارية..."):
                
                full_user_idea = f"""
                - اسم الابتكار ووظيفته: {input_1}
                - الآلية التقنية المقترحة: {input_2}
                - المشكلة والفجوة المستهدفة: {input_3}
                """
                
                prompt = f"""
                بصفتك خبير براءات اختراع عالمي ومستشاراً استراتيجياً لابتكار القيمة وفرض الريادة الابتكارية، حلل معطيات الابتكار الموزعة على الخانات التالية (والصورة المرفقة إن وجدت):
                {full_user_idea}

                ⚠️ قواعد صارمة وحاسمة للتقييم (بروتوكول الحزم):
                إذا تبين لك أن المعطيات المدخلة في الخانات أعلاه شحيحة جداً، أو مكررة، أو عامة وخالية من "خلطة سرية أو آلية تقنية واضحة ومحددة"، أو أن المستخدم يتهرب من الإجابة الفنية، التزم تماماً بالآتي:
                1. اصدر حكماً تاديبياً صارماً، واجعل التقييم النهائي الإجمالي في وسم [===SCORE===] منخفضاً جداً (بين 10 إلى 45 كحد أقصى) ليعكس واقعية ضعف الطرح التقني.
                2. في قسم التشخيص، لا تقم باختراع أو تأليف أي تفاصيل هندسية أو حساسات من عندك نيابة عن المستخدم، بل واجهه بوضوح واكتب له: (بناءً على معطياتك الشحيحة، فكرتك مكشوفة تماماً ولا ترتقي لتسجيل براءة اختراع للأسباب التالية...).
                3. في قسم خطة الريادة، تحوّل إلى موجه هندسي حازم، واكتب له "بروتوكول سد الفجوة التقنية" على شكل أسئلة هندسية وقانونية محددة يجب عليه البحث عنها وإجابتها ليرفع من قيمة فكرته.

                أما إذا كانت المعطيات غنية وجادة ومتكاملة، فقم بإنتاج التقرير الاحترافي الكامل المعتاد.
                التزم بالأوسمة التالية حصراً لتقسيم التقرير، واجعل الفراغات واضحة:

                [===LEVEL1===] 
                ### 📊 أولاً: التشخيص الاستراتيجي والجوهر الهندسي [cite: 3]
                1. تفكيك المنظومة التقنية الحالية (أو نقد المعطيات إن كانت ناقصة). [cite: 4]
                2. السيناريو التشغيلي والديناميكي على أرض الواقع. [cite: 9]

                [===COMPETITORS_TABLE===]
                صغ جدول مقارنة بالمؤشرات (Markdown Table) يقارن هذه المعطيات مع أبرز 3 منافسين بناءً على الأعمدة التالية حصراً: (المنافس | الفجوة التقنية لديه | ميزتك التنافسية الصارمة أو الفجوة التي يجب عليك سدها لتنافسه). لا تكتب أي نص خارج حدود الجدول في هذا القسم. [cite: 15]

                [===LEVEL2===]
                ### 🔧 ثانياً: المطالبات التقنية والمراجع القانونية [cite: 23]
                1. صياغة مسودة المطالبات التقنية (Patent Claims) بناءً على ما قدمه (أو توضيح استحالة الصياغة لغياب الآلية). [cite: 24]
                2. الفن السابق (Prior Art): الثغرات في براءات الاختراع الحالية وكيف تتجاوزها هذه الفكرة. [cite: 28]

                [===LEVEL3===]
                ### 🛣️ ثالثاً: الجدوى الاقتصادية وخارطة الريادة الابتكارية [cite: 32]
                1. نموذج فرض الريادة وحماية الهندسة العكسية. [cite: 33]
                2. خارطة الطريق التنفيذية وسد الفجوة المعرفية. [cite: 37]

                [===AUDIT===]
                ### ⭐ رابعاً: تفصيل مؤشر الفرادة [cite: 42]
                قم بتقييم الفكرة من 100 بناءً على الأوزان التالية مع تبرير النسبة الممنوحة باختصار شديد: [cite: 42]
                1. الجدة وعدم وجود فن سابق (30%) [cite: 43]
                2. الخطوة الابتكارية وعدم البداهة (25%) [cite: 44]
                3. تميز البنية الهندسية وتدفق البيانات (15%) [cite: 45]
                4. القابلية للتطبيق وحل المشكلة التقنية (15%) [cite: 46]
                5. صعوبة الهندسة العكسية وبناء الخندق المائي (15%) [cite: 47]

                [===SCORE===]
                اكتب الرقم الإجمالي النهائي فقط لنسبة الفرادة الموزونة بين الأوسمة (رقم مجرد فقط)، مثال: 75 [cite: 48]
                [===/SCORE===]

                [===SOVEREIGNTY===]
                ### 💡 خامساً: توصية الريادة الابتكارية المطلقة [cite: 49]
                قدم التوصية الاستراتيجية الذهبية للمبتكر لفرض ريادته التجارية بناءً على حالته (تطويرية أو تصحيحية حازمة).
                """
                
                st.session_state.full_report = call_pro_api(prompt, uploaded_file)
                st.session_state.final_idea = input_1 
                st.session_state.gate_passed = True
                st.rerun()
        else:
            st.error("⚠️ من فضلك، يجب ملء الخانات الثلاث الإلزامية بالكامل لتفعيل فحص الرادار!")
else:
    # المرحلة الثانية: العرض البصري الآمن والمحمي 100%
    report = st.session_state.full_report
    
    score_match = re.search(r'\[===SCORE===\]\s*(\d+)\s*\[===/SCORE===\]', report)
    innovation_score = int(score_match.group(1)) if score_match else 75
        
# تنظيف النص المعروض من أوسمة السكور
    clean_report = re.sub(r'\[===SCORE===\].*?\[===/SCORE===\]', '', report, flags=re.DOTALL)
    
    # تحصين تفكيك أجزاء البروتوكول الحازم لمنع الـ NameError نهائياً
    parts = re.split(r'\[===LEVEL[1-3]===\]|\[===COMPETITORS_TABLE===\]|\[===AUDIT===\]|\[===SOVEREIGNTY===\]', clean_report)
    
    # تعريف افتراضي لجميع الأقسام لتفادي نقص العناصر في المصفوفة
    level1_text = ""
    competitors_table = ""
    level2_text = ""
    level3_text = ""
    audit_text = ""
    sovereignty_text = ""

    # تعبئة المتغيرات بناءً على ما رجع من المحرك فعلياً
    if len(parts) > 1: level1_text = parts[1]
    if len(parts) > 2: competitors_table = parts[2]
    if len(parts) > 3: level2_text = parts[3]
    if len(parts) > 4: level3_text = parts[4]
    if len(parts) > 5: audit_text = parts[5]
    if len(parts) > 6: sovereignty_text = parts[6]

    # إذا فشل التفكيك تماماً، نضع النص كاملاً في التبويب الأول كخطة احتياطية
    if not level1_text.strip():
        level1_text = clean_report

    st.markdown("### 🎯 التقييم الفوري ومستوى جدية الابتكار")
    
    with tab1:
        st.markdown(f"<div style='direction:rtl; text-align:right;'>{level1_text}</div>", unsafe_allow_html=True)
        if competitors_table.strip():
            st.markdown("#### 🔄 مصفوفة المقارنة التنافسية وتحديد الفجوات")
            st.markdown(f"<div style='direction:rtl; text-align:right;'>{competitors_table}</div>", unsafe_allow_html=True) [cite: 15]
            
            st.divider()
            # 🔥 زر استعلام جوجل باتنت الحي المباشر دون الحاجة لحسابات
            search_query = st.session_state.final_idea.replace(" ", "+")
            google_patents_url = f"https://patents.google.com/?q={search_query}&sort=new"
            st.markdown(f"""
                <div style="text-align: right; direction: rtl; margin-top: 15px;">
                    <p style="font-size: 12pt; color: #475569;">💡 <b>فحص حي إضافي:</b> يمكنك التحقق من براءات الاختراع المودعة عالمياً والمشابهة لابتكارك مباشرة عبر الرابط الديناميكي التالي:</p>
                    <a href="{google_patents_url}" target="_blank" style="background-color: #1e3a8a; color: white; padding: 10px 20px; text-decoration: none; border-radius: 5px; font-weight: bold; display: inline-block;">🔍 استكشاف الفكرة حياً في Google Patents</a>
                </div>
            """, unsafe_allow_html=True)
            
    with tab2:
        st.markdown(f"<div style='direction:rtl; text-align:right;'>{level2_text}</div>", unsafe_allow_html=True)
        st.divider()
        
        # 🔥 حقن مرجعية تشريعات اليمن لـ WIPO Lex في قلب تبويب القوانين
        wipo_ye_url = "https://www.wipo.int/wipolex/ar/legislation/results?countryOrgs=YE&last=true&activeCollection=laws&collection=laws&collection=treaties&collection=judgments"
        st.markdown(f"""
            <div style="background-color: #f1f5f9; padding: 15px; border-radius: 6px; border-right: 4px solid #475569; text-align: right; direction: rtl; margin-bottom: 20px;">
                <b>⚖️ المرجعية القانونية والتشريعية المحلية (WIPO Lex):</b><br>
                لتكييف المطالبات القانونية أعلاه وحمايتها محلياً، تم ربط الفحص بالقوانين والاتفاقيات النافذة في الجمهورية اليمنية المعتمدة لدى المنظمة العالمية للملكية الفكرية. 
                <br><a href="{wipo_ye_url}" target="_blank" style="color: #1e3a8a; font-weight: bold; text-decoration: underline; display: inline-block; margin-top: 8px;">🔗 اضغط هنا لتصفح جدول التشريعات والمعاهدات اليمنية النافذة في WIPO</a>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"<div style='direction:rtl; text-align:right;'>{audit_text}</div>", unsafe_allow_html=True)
        
    with tab3:
        st.markdown(f"<div style='direction:rtl; text-align:right;'>{level3_text}</div>", unsafe_allow_html=True)

    st.divider()
    
    # عرض التوصية الكبرى المعدلة بمسماها الجديد
    st.warning(f"**💡 توصية الريادة الابتكارية (مستشارك الرقمي)**\n\n{sovereignty_text.replace('### خامساً: توصية السيادة التسويقية المطلقة', '')}") [cite: 49]

    st.divider()
    
    docx_file = create_docx(report, st.session_state.final_idea)
    if docx_file:
        st.download_button(
            label="📥 تحميل التقرير الاستشاري كملف Word منسق موسط",
            data=docx_file,
            file_name="Innovation_Sovereignty_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    if st.button("🔄 فحص ابتكار جديد"):
        st.session_state.clear()
        st.rerun()
